"""
Ingestion service for the RAG module.

This module provides functionality for ingesting documents from various sources.
"""

import uuid
import json
import asyncio
from datetime import datetime
from typing import Dict, Any, List, Optional, Tuple, Union
import httpx
import io
import tempfile
import os

from app.core.logger import logger
from app.core.database import get_db_client
from app.rag.services.chunking_service import ChunkingService
from app.rag.services.embedding_service import EmbeddingService
from app.rag.repositories.processing_job_repository import ProcessingJobRepository
from app.rag.repositories.document_repository import DocumentRepository


class IngestionCoordinator:
    """Coordinator for document ingestion."""

    def __init__(
        self,
        chunking_service: ChunkingService = None,
        embedding_service: EmbeddingService = None,
        processing_job_repository: ProcessingJobRepository = None,
        document_repository: DocumentRepository = None
    ):
        """Initialize the ingestion coordinator.
        
        Args:
            chunking_service: Service for chunking text.
            embedding_service: Service for generating embeddings.
            processing_job_repository: Repository for processing jobs.
            document_repository: Repository for documents.
        """
        self.chunking_service = chunking_service or ChunkingService()
        self.embedding_service = embedding_service or EmbeddingService()
        self.processing_job_repository = processing_job_repository or ProcessingJobRepository()
        self.document_repository = document_repository or DocumentRepository()
        
        # Initialize processors
        self.processors = {
            "file": FileProcessor(),
            "url": URLProcessor(),
            "text": TextProcessor()
        }

    async def process_document(
        self,
        source_type: str,
        collection_id: str,
        content: Union[bytes, str],
        metadata: Dict[str, Any]
    ) -> str:
        """Process a document from a source.
        
        Args:
            source_type: Type of source (file, url, text).
            collection_id: ID of the collection to add the document to.
            content: Content of the document.
            metadata: Metadata for the document.
            
        Returns:
            ID of the processing job.
            
        Raises:
            ValueError: If the source type is not supported.
        """
        if source_type not in self.processors:
            raise ValueError(f"Unsupported source type: {source_type}")
        
        # Create document
        document_id = str(uuid.uuid4())
        now = datetime.utcnow()
        
        document_data = {
            "id": document_id,
            "collection_id": collection_id,
            "name": metadata.get("name", "Untitled Document"),
            "description": metadata.get("description"),
            "source_type": source_type,
            "source_url": metadata.get("url") if source_type == "url" else None,
            "file_type": metadata.get("file_type") if source_type == "file" else None,
            "file_size": metadata.get("file_size") if source_type == "file" else None,
            "status": "processing",
            "created_at": now,
            "updated_at": now
        }
        
        # Create document in database
        await self.document_repository.create(document_data)
        
        # Create processing job
        job_id = str(uuid.uuid4())
        job_data = {
            "id": job_id,
            "document_id": document_id,
            "status": "created",
            "progress": 0.0,
            "created_at": now,
            "updated_at": now
        }
        
        # Create processing job in database
        await self.processing_job_repository.create(job_data)
        
        # Start processing in background
        asyncio.create_task(self._process_document_async(
            job_id=job_id,
            document_id=document_id,
            source_type=source_type,
            content=content,
            metadata=metadata
        ))
        
        return job_id

    async def _process_document_async(
        self,
        job_id: str,
        document_id: str,
        source_type: str,
        content: Union[bytes, str],
        metadata: Dict[str, Any]
    ) -> None:
        """Process a document asynchronously.
        
        Args:
            job_id: ID of the processing job.
            document_id: ID of the document.
            source_type: Type of source (file, url, text).
            content: Content of the document.
            metadata: Metadata for the document.
        """
        try:
            # Update job status
            await self.processing_job_repository.update_status(
                job_id=job_id,
                status="processing",
                progress=0.1
            )
            
            # Extract text from source
            processor = self.processors[source_type]
            text, extracted_metadata = await processor.extract_text(content, metadata)
            
            # Update job progress
            await self.processing_job_repository.update_status(
                job_id=job_id,
                status="processing",
                progress=0.3
            )
            
            # Chunk text
            chunk_size = metadata.get("chunk_size", 1000)
            chunk_overlap = metadata.get("chunk_overlap", 200)
            chunking_strategy = metadata.get("chunking_strategy", "fixed_size")
            
            chunks = self.chunking_service.chunk_text(
                text=text,
                chunk_size=chunk_size,
                chunk_overlap=chunk_overlap,
                chunking_strategy=chunking_strategy,
                metadata=extracted_metadata
            )
            
            # Update job progress
            await self.processing_job_repository.update_status(
                job_id=job_id,
                status="processing",
                progress=0.5
            )
            
            # Generate embeddings
            texts = [chunk["content"] for chunk in chunks]
            embeddings = await self.embedding_service.generate_embeddings(texts)
            
            # Update job progress
            await self.processing_job_repository.update_status(
                job_id=job_id,
                status="processing",
                progress=0.7
            )
            
            # Store chunks and embeddings
            chunk_ids = []
            for i, (chunk, embedding) in enumerate(zip(chunks, embeddings)):
                chunk_id = str(uuid.uuid4())
                chunk_ids.append(chunk_id)
                
                # Store chunk in database
                await self.document_repository.create_chunk(
                    chunk_id=chunk_id,
                    document_id=document_id,
                    content=chunk["content"],
                    chunk_index=chunk["chunk_index"],
                    metadata=chunk["metadata"],
                    embedding=embedding
                )
            
            # Update job progress
            await self.processing_job_repository.update_status(
                job_id=job_id,
                status="processing",
                progress=0.9
            )
            
            # Update document status
            await self.document_repository.update_status(
                document_id=document_id,
                status="processed"
            )
            
            # Update job status
            await self.processing_job_repository.update_status(
                job_id=job_id,
                status="completed",
                progress=1.0
            )
            
        except Exception as e:
            logger.error(f"Error processing document: {str(e)}")
            
            # Update document status
            await self.document_repository.update_status(
                document_id=document_id,
                status="failed"
            )
            
            # Update job status
            await self.processing_job_repository.update_status(
                job_id=job_id,
                status="failed",
                progress=0.0,
                error_message=str(e)
            )

    async def get_processing_status(self, job_id: str) -> Dict[str, Any]:
        """Get the status of a processing job.
        
        Args:
            job_id: ID of the processing job.
            
        Returns:
            Status of the processing job.
        """
        return await self.processing_job_repository.get_by_id(job_id)


class FileProcessor:
    """Processor for file sources."""

    async def extract_text(
        self,
        content: bytes,
        metadata: Dict[str, Any]
    ) -> Tuple[str, Dict[str, Any]]:
        """Extract text from a file.
        
        Args:
            content: Content of the file.
            metadata: Metadata for the file.
            
        Returns:
            Tuple of extracted text and metadata.
            
        Raises:
            ValueError: If the file type is not supported.
        """
        file_type = metadata.get("file_type", "")
        
        if "pdf" in file_type.lower():
            return await self._extract_text_from_pdf(content, metadata)
        elif "word" in file_type.lower() or "docx" in file_type.lower():
            return await self._extract_text_from_docx(content, metadata)
        elif "text" in file_type.lower() or "txt" in file_type.lower():
            return await self._extract_text_from_txt(content, metadata)
        elif "csv" in file_type.lower():
            return await self._extract_text_from_csv(content, metadata)
        elif "json" in file_type.lower():
            return await self._extract_text_from_json(content, metadata)
        elif "html" in file_type.lower() or "htm" in file_type.lower():
            return await self._extract_text_from_html(content, metadata)
        else:
            raise ValueError(f"Unsupported file type: {file_type}")

    async def _extract_text_from_pdf(
        self,
        content: bytes,
        metadata: Dict[str, Any]
    ) -> Tuple[str, Dict[str, Any]]:
        """Extract text from a PDF file.
        
        Args:
            content: Content of the PDF file.
            metadata: Metadata for the PDF file.
            
        Returns:
            Tuple of extracted text and metadata.
        """
        try:
            import pypdf
            
            # Create a temporary file
            with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as temp_file:
                temp_file.write(content)
                temp_file_path = temp_file.name
            
            try:
                # Extract text from PDF
                pdf_reader = pypdf.PdfReader(temp_file_path)
                
                # Extract metadata
                pdf_metadata = pdf_reader.metadata
                extracted_metadata = {
                    "title": pdf_metadata.get("/Title", metadata.get("name", "")),
                    "author": pdf_metadata.get("/Author", ""),
                    "subject": pdf_metadata.get("/Subject", ""),
                    "creator": pdf_metadata.get("/Creator", ""),
                    "producer": pdf_metadata.get("/Producer", ""),
                    "page_count": len(pdf_reader.pages)
                }
                
                # Extract text
                text = ""
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n\n"
                
                return text.strip(), extracted_metadata
            finally:
                # Clean up temporary file
                os.unlink(temp_file_path)
                
        except ImportError:
            logger.warning("pypdf not installed. Using fallback method for PDF extraction.")
            return content.decode("utf-8", errors="ignore"), metadata
        except Exception as e:
            logger.error(f"Error extracting text from PDF: {str(e)}")
            return content.decode("utf-8", errors="ignore"), metadata

    async def _extract_text_from_docx(
        self,
        content: bytes,
        metadata: Dict[str, Any]
    ) -> Tuple[str, Dict[str, Any]]:
        """Extract text from a DOCX file.
        
        Args:
            content: Content of the DOCX file.
            metadata: Metadata for the DOCX file.
            
        Returns:
            Tuple of extracted text and metadata.
        """
        try:
            import docx
            
            # Create a temporary file
            with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as temp_file:
                temp_file.write(content)
                temp_file_path = temp_file.name
            
            try:
                # Extract text from DOCX
                doc = docx.Document(temp_file_path)
                
                # Extract metadata
                core_properties = doc.core_properties
                extracted_metadata = {
                    "title": core_properties.title or metadata.get("name", ""),
                    "author": core_properties.author or "",
                    "subject": core_properties.subject or "",
                    "paragraph_count": len(doc.paragraphs)
                }
                
                # Extract text
                text = ""
                for paragraph in doc.paragraphs:
                    text += paragraph.text + "\n\n"
                
                return text.strip(), extracted_metadata
            finally:
                # Clean up temporary file
                os.unlink(temp_file_path)
                
        except ImportError:
            logger.warning("python-docx not installed. Using fallback method for DOCX extraction.")
            return content.decode("utf-8", errors="ignore"), metadata
        except Exception as e:
            logger.error(f"Error extracting text from DOCX: {str(e)}")
            return content.decode("utf-8", errors="ignore"), metadata

    async def _extract_text_from_txt(
        self,
        content: bytes,
        metadata: Dict[str, Any]
    ) -> Tuple[str, Dict[str, Any]]:
        """Extract text from a TXT file.
        
        Args:
            content: Content of the TXT file.
            metadata: Metadata for the TXT file.
            
        Returns:
            Tuple of extracted text and metadata.
        """
        try:
            text = content.decode("utf-8")
            
            # Extract metadata
            extracted_metadata = {
                "title": metadata.get("name", ""),
                "line_count": text.count("\n") + 1,
                "char_count": len(text)
            }
            
            return text, extracted_metadata
        except Exception as e:
            logger.error(f"Error extracting text from TXT: {str(e)}")
            return content.decode("utf-8", errors="ignore"), metadata

    async def _extract_text_from_csv(
        self,
        content: bytes,
        metadata: Dict[str, Any]
    ) -> Tuple[str, Dict[str, Any]]:
        """Extract text from a CSV file.
        
        Args:
            content: Content of the CSV file.
            metadata: Metadata for the CSV file.
            
        Returns:
            Tuple of extracted text and metadata.
        """
        try:
            import csv
            import io
            
            # Decode content
            text_io = io.StringIO(content.decode("utf-8"))
            
            # Read CSV
            reader = csv.reader(text_io)
            rows = list(reader)
            
            # Extract metadata
            extracted_metadata = {
                "title": metadata.get("name", ""),
                "row_count": len(rows),
                "column_count": len(rows[0]) if rows else 0
            }
            
            # Convert to text
            text = ""
            for row in rows:
                text += " | ".join(row) + "\n"
            
            return text, extracted_metadata
        except Exception as e:
            logger.error(f"Error extracting text from CSV: {str(e)}")
            return content.decode("utf-8", errors="ignore"), metadata

    async def _extract_text_from_json(
        self,
        content: bytes,
        metadata: Dict[str, Any]
    ) -> Tuple[str, Dict[str, Any]]:
        """Extract text from a JSON file.
        
        Args:
            content: Content of the JSON file.
            metadata: Metadata for the JSON file.
            
        Returns:
            Tuple of extracted text and metadata.
        """
        try:
            # Decode content
            json_str = content.decode("utf-8")
            
            # Parse JSON
            json_data = json.loads(json_str)
            
            # Extract metadata
            extracted_metadata = {
                "title": metadata.get("name", ""),
                "json_type": type(json_data).__name__
            }
            
            # Convert to text
            text = json.dumps(json_data, indent=2)
            
            return text, extracted_metadata
        except Exception as e:
            logger.error(f"Error extracting text from JSON: {str(e)}")
            return content.decode("utf-8", errors="ignore"), metadata

    async def _extract_text_from_html(
        self,
        content: bytes,
        metadata: Dict[str, Any]
    ) -> Tuple[str, Dict[str, Any]]:
        """Extract text from an HTML file.
        
        Args:
            content: Content of the HTML file.
            metadata: Metadata for the HTML file.
            
        Returns:
            Tuple of extracted text and metadata.
        """
        try:
            from bs4 import BeautifulSoup
            
            # Decode content
            html_str = content.decode("utf-8", errors="ignore")
            
            # Parse HTML
            soup = BeautifulSoup(html_str, "html.parser")
            
            # Extract metadata
            title = soup.title.string if soup.title else metadata.get("name", "")
            extracted_metadata = {
                "title": title,
                "links": len(soup.find_all("a")),
                "images": len(soup.find_all("img"))
            }
            
            # Extract text
            for script in soup(["script", "style"]):
                script.extract()
            
            text = soup.get_text(separator="\n")
            lines = (line.strip() for line in text.splitlines())
            chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
            text = "\n".join(chunk for chunk in chunks if chunk)
            
            return text, extracted_metadata
        except ImportError:
            logger.warning("beautifulsoup4 not installed. Using fallback method for HTML extraction.")
            return content.decode("utf-8", errors="ignore"), metadata
        except Exception as e:
            logger.error(f"Error extracting text from HTML: {str(e)}")
            return content.decode("utf-8", errors="ignore"), metadata


class URLProcessor:
    """Processor for URL sources."""

    async def extract_text(
        self,
        url: str,
        metadata: Dict[str, Any]
    ) -> Tuple[str, Dict[str, Any]]:
        """Extract text from a URL.
        
        Args:
            url: URL to extract text from.
            metadata: Metadata for the URL.
            
        Returns:
            Tuple of extracted text and metadata.
        """
        try:
            # Check if we should use Firecrawl
            use_firecrawl = metadata.get("use_firecrawl", True)
            
            if use_firecrawl:
                return await self._extract_text_with_firecrawl(url, metadata)
            else:
                return await self._extract_text_with_httpx(url, metadata)
        except Exception as e:
            logger.error(f"Error extracting text from URL: {str(e)}")
            return f"Failed to extract text from URL: {url}", metadata

    async def _extract_text_with_firecrawl(
        self,
        url: str,
        metadata: Dict[str, Any]
    ) -> Tuple[str, Dict[str, Any]]:
        """Extract text from a URL using Firecrawl.
        
        Args:
            url: URL to extract text from.
            metadata: Metadata for the URL.
            
        Returns:
            Tuple of extracted text and metadata.
        """
        from app.core.config import get_settings
        
        firecrawl_api_key = get_settings().firecrawl_api_key
        firecrawl_url = get_settings().firecrawl_url or "https://firecrawl.dev"
        
        if not firecrawl_api_key:
            logger.warning("Firecrawl API key not set. Falling back to HTTPX.")
            return await self._extract_text_with_httpx(url, metadata)
        
        try:
            # Make request to Firecrawl
            async with httpx.AsyncClient() as client:
                response = await client.post(
                    f"{firecrawl_url}/api/v1/crawl",
                    headers={
                        "Authorization": f"Bearer {firecrawl_api_key}",
                        "Content-Type": "application/json"
                    },
                    json={
                        "url": url,
                        "javascript": True,
                        "wait_for": 2000
                    },
                    timeout=30.0
                )
                
                response.raise_for_status()
                result = response.json()
                
                # Extract text
                text = result.get("text", "")
                
                # Extract metadata
                extracted_metadata = {
                    "title": result.get("title", metadata.get("name", url)),
                    "url": url,
                    "links": len(result.get("links", [])),
                    "status_code": result.get("status_code"),
                    "content_type": result.get("content_type")
                }
                
                return text, extracted_metadata
        except Exception as e:
            logger.error(f"Error extracting text with Firecrawl: {str(e)}")
            return await self._extract_text_with_httpx(url, metadata)

    async def _extract_text_with_httpx(
        self,
        url: str,
        metadata: Dict[str, Any]
    ) -> Tuple[str, Dict[str, Any]]:
        """Extract text from a URL using HTTPX.
        
        Args:
            url: URL to extract text from.
            metadata: Metadata for the URL.
            
        Returns:
            Tuple of extracted text and metadata.
        """
        try:
            # Make request to URL
            async with httpx.AsyncClient() as client:
                response = await client.get(url, follow_redirects=True, timeout=10.0)
                response.raise_for_status()
                
                # Get content type
                content_type = response.headers.get("content-type", "").lower()
                
                # Process based on content type
                if "text/html" in content_type:
                    # Process as HTML
                    file_processor = FileProcessor()
                    return await file_processor._extract_text_from_html(
                        response.content,
                        {**metadata, "name": metadata.get("name", url)}
                    )
                elif "application/pdf" in content_type:
                    # Process as PDF
                    file_processor = FileProcessor()
                    return await file_processor._extract_text_from_pdf(
                        response.content,
                        {**metadata, "name": metadata.get("name", url)}
                    )
                elif "application/json" in content_type:
                    # Process as JSON
                    file_processor = FileProcessor()
                    return await file_processor._extract_text_from_json(
                        response.content,
                        {**metadata, "name": metadata.get("name", url)}
                    )
                else:
                    # Process as plain text
                    text = response.text
                    
                    # Extract metadata
                    extracted_metadata = {
                        "title": metadata.get("name", url),
                        "url": url,
                        "content_type": content_type,
                        "status_code": response.status_code
                    }
                    
                    return text, extracted_metadata
        except Exception as e:
            logger.error(f"Error extracting text with HTTPX: {str(e)}")
            return f"Failed to extract text from URL: {url}", metadata


class TextProcessor:
    """Processor for text sources."""

    async def extract_text(
        self,
        text: str,
        metadata: Dict[str, Any]
    ) -> Tuple[str, Dict[str, Any]]:
        """Extract text from raw text.
        
        Args:
            text: Raw text.
            metadata: Metadata for the text.
            
        Returns:
            Tuple of extracted text and metadata.
        """
        # Extract metadata
        extracted_metadata = {
            "title": metadata.get("name", "Text Document"),
            "line_count": text.count("\n") + 1,
            "char_count": len(text)
        }
        
        return text, extracted_metadata